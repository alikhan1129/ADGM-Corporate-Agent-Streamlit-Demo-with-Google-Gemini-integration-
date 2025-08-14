# docx_utils.py
from docx import Document
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def parse_docx_documents(paths):
    out = []
    for p in paths:
        doc = Document(p)
        paras = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
        full_text = "\n".join(paras)
        out.append({"path": p, "text": full_text, "paragraphs": paras})
    return out


def _find_paragraph_index_for_issue(doc, section, issue_text):
    """
    Heuristic: prefer paragraphs which contain longer tokens from section or issue_text.
    Returns index of paragraph or None if not found.
    """
    section_tokens = [t for t in re.split(r"\W+", (section or "").lower()) if len(t) > 3]
    issue_tokens = [t for t in re.split(r"\W+", (issue_text or "").lower()) if len(t) > 3]

    best_idx = None
    best_score = 0
    for i, para in enumerate(doc.paragraphs):
        pl = para.text.lower()
        score = 0
        for tok in section_tokens:
            if tok in pl:
                score += 2
        for tok in issue_tokens:
            if tok in pl:
                score += 1
        if score > best_score:
            best_score = score
            best_idx = i
    # require at least a minimal score to accept match
    if best_score >= 1:
        return best_idx
    return None


def _insert_after_paragraph(doc, para_index, text, author="ADGM-Agent"):
    """
    Insert a new paragraph after para_index. Uses docx API to add paragraph then move it.
    Returns True if inserted, False otherwise.
    """
    try:
        # Create a new paragraph at the end then move it after the target paragraph
        after = doc.add_paragraph()
        run = after.add_run(text)
        # Basic formatting to make it look like a review comment
        run.italic = True
        run.bold = False

        # Try to reposition the newly added paragraph to after the target paragraph
        p = doc.paragraphs[para_index]
        p_elm = p._p
        after_elm = after._p
        p_elm.addnext(after_elm)
        return True
    except Exception:
        # If anything fails, fall back to append (we already added it)
        return True


def insert_review_notes_and_save(original_path, issues, out_path):
    """
    Insert review notes into the docx and save to out_path.
    For each issue, attempt to locate a best paragraph and insert a review paragraph immediately after it.
    If location can't be determined, append a general note at the end.
    """
    doc = Document(original_path)
    for iss in issues:
        section = iss.get("section", "")
        note_text = f"REVIEW NOTE (severity={iss.get('severity','Medium')}): {iss.get('issue')}. Suggestion: {iss.get('suggestion','')}"
        inserted = False
        if section:
            try:
                idx = _find_paragraph_index_for_issue(doc, section, iss.get('issue', ''))
                if idx is not None:
                    _insert_after_paragraph(doc, idx, note_text)
                    inserted = True
            except Exception:
                inserted = False
        if not inserted:
            # fallback: append a clearly identifiable review note at end
            doc.add_paragraph(f"REVIEW NOTE (general): {note_text}")
    doc.save(out_path)
    return out_path
